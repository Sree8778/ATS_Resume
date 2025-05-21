import streamlit as st
import docx2txt
from io import StringIO
import base64

st.set_page_config(page_title="ATS Resume Converter", layout="centered")

# --- Custom Styles ---
st.markdown("""import streamlit as st
import docx2txt
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO

st.set_page_config(page_title="ATS Resume Formatter", layout="centered", page_icon="📄")

st.markdown("""
    <style>
    .block-container { padding: 2rem 2rem 2rem 2rem; }
    .stButton>button {
        font-size: 16px;
        padding: 0.5em 2em;
        background-color: #2c7be5;
        color: white;
        border: none;
        border-radius: 6px;
    }
    </style>
""", unsafe_allow_html=True)

st.title("📄 ATS Resume Formatter")
st.write("Upload a `.docx` or `.pdf` resume and get an ATS-optimized version instantly.")

uploaded_file = st.file_uploader("📂 Upload Resume", type=["docx", "pdf"])

# Extract text based on file type
def extract_text(file, file_type):
    if file_type == "docx":
        return docx2txt.process(file)
    elif file_type == "pdf":
        pdf = fitz.open(stream=file.read(), filetype="pdf")
        text = ""
        for page in pdf:
            text += page.get_text()
        return text
    return ""

# Convert to ATS format
def convert_to_ats_format(text):
    sections = {
        "Contact": "", "Education": "", "Experience": "",
        "Skills": "", "Projects": "", "Certifications": ""
    }
    current_section = "Contact"
    for line in text.splitlines():
        line = line.strip()
        if not line: continue
        tag = line.lower()
        if "education" in tag:
            current_section = "Education"
        elif "experience" in tag or "work" in tag:
            current_section = "Experience"
        elif "skill" in tag:
            current_section = "Skills"
        elif "project" in tag:
            current_section = "Projects"
        elif "certification" in tag:
            current_section = "Certifications"
        else:
            sections[current_section] += line + "\n"
    return sections

# Export as .docx
def export_docx(sections):
    doc = Document()
    for section, content in sections.items():
        doc.add_heading(section.upper(), level=2)
        for line in content.splitlines():
            doc.add_paragraph(line)
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

if uploaded_file:
    file_type = uploaded_file.name.split(".")[-1]
    text = extract_text(uploaded_file, file_type)
    structured = convert_to_ats_format(text)

    st.success("✅ Resume parsed and converted successfully!")
    st.markdown("### 📝 Preview ATS Resume")
    for sec, val in structured.items():
        st.markdown(f"**{sec}**")
        st.text(val)

    ats_txt = "".join(f"===== {k.upper()} =====\n{v}\n" for k, v in structured.items())
    ats_docx = export_docx(structured)

    col1, col2 = st.columns(2)
    with col1:
        st.download_button("📥 Download .txt", data=ats_txt, file_name="ATS_Resume.txt")
    with col2:
        st.download_button("📥 Download .docx", data=ats_docx, file_name="ATS_Resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    st.info("👆 Drag and drop a `.docx` or `.pdf` resume to begin.")

    <style>
    .main {
        background-color: #f5f7fa;
    }
    .stTextArea, .stTextInput, .stFileUploader, .stButton {
        font-size: 16px !important;
    }
    .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }
    </style>
""", unsafe_allow_html=True)

# --- Title & Instructions ---
st.title("🚀 ATS Resume Formatter")
st.subheader("Convert any resume into an ATS-compliant format with one click.")
st.markdown("💼 Upload your **.docx** resume using the drag & drop area below.")

# --- File Uploader ---
uploaded_file = st.file_uploader("📄 Upload Resume File (.docx)", type=["docx"], label_visibility="collapsed")

# --- Resume Processing Functions ---
def extract_text_from_docx(file):
    return docx2txt.process(file)

def convert_to_ats_format(text):
    ats_sections = {
        "Contact": "",
        "Education": "",
        "Experience": "",
        "Skills": "",
        "Projects": "",
        "Certifications": ""
    }

    lines = text.split('\n')
    current_section = "Contact"

    for line in lines:
        line = line.strip()
        if not line:
            continue

        header = line.lower()
        if "education" in header:
            current_section = "Education"
        elif "experience" in header or "work" in header:
            current_section = "Experience"
        elif "skill" in header:
            current_section = "Skills"
        elif "project" in header:
            current_section = "Projects"
        elif "certification" in header:
            current_section = "Certifications"
        else:
            ats_sections[current_section] += line + "\n"

    return ats_sections

# --- If File is Uploaded ---
if uploaded_file:
    st.success("✅ Resume uploaded successfully!")

    text = extract_text_from_docx(uploaded_file)
    ats_resume = convert_to_ats_format(text)

    st.markdown("### 📑 Preview of ATS-Formatted Resume")

    output_text = ""
    for section, content in ats_resume.items():
        st.markdown(f"#### ✦ {section}")
        st.text(content)
        output_text += f"===== {section.upper()} =====\n{content}\n"

    # --- Download Option ---
    st.download_button(
        label="⬇️ Download Clean ATS Resume (.txt)",
        data=output_text,
        file_name="ATS_Resume.txt",
        mime="text/plain"
    )

else:
    st.info("👆 Drag and drop a .docx resume file to begin.")
